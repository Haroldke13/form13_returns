#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import load_workbook


@dataclass
class ChildOrganization:
    pbo_name: str
    projects: str
    scope: str
    counties: str
    report_ids: str
    report_count: int | None
    project_count: int | None


@dataclass
class ContactDetails:
    name: str
    phone: str
    email: str


def normalize_space(value: object) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def normalize_key(value: object) -> str:
    return re.sub(r"[^A-Z0-9]+", " ", normalize_space(value).upper()).strip()


def child_related(text: object) -> bool:
    return "CHILD" in normalize_space(text).upper()


def coverage_includes_county(organization: ChildOrganization, county: str) -> bool:
    return normalize_space(county).upper() in organization.counties.upper()


def first_nonempty(values: Iterable[object]) -> str:
    for value in values:
        cleaned = normalize_space(value)
        if cleaned:
            return cleaned
    return ""


def identify_child_organizations(projects_path: Path) -> dict[str, ChildOrganization]:
    workbook = load_workbook(projects_path, read_only=True, data_only=True)
    sheet = workbook["PBO Summary"]
    headers = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
    header_map = {str(name): index for index, name in enumerate(headers)}

    organizations: dict[str, ChildOrganization] = {}
    for row in sheet.iter_rows(min_row=2, values_only=True):
        pbo_name = normalize_space(row[header_map["pbo_name"]])
        projects = normalize_space(row[header_map["projects"]])
        if not pbo_name:
            continue
        if not (child_related(pbo_name) or child_related(projects)):
            continue

        key = normalize_key(pbo_name)
        organizations[key] = ChildOrganization(
            pbo_name=pbo_name,
            projects=projects,
            scope=normalize_space(row[header_map["scope"]]),
            counties=normalize_space(row[header_map["report_counties"]]),
            report_ids=normalize_space(row[header_map["report_ids"]]),
            report_count=row[header_map["report_count"]],
            project_count=row[header_map["project_count"]],
        )

    return organizations


def merge_contact_rows(rows: list[dict[str, object]]) -> ContactDetails:
    return ContactDetails(
        name=first_nonempty(row.get("contact_name") for row in rows),
        phone=first_nonempty(
            field
            for row in rows
            for field in (
                row.get("contact_telephone"),
                row.get("telephone"),
                row.get("cell_phone"),
            )
        ),
        email=first_nonempty(
            field
            for row in rows
            for field in (
                row.get("contact_email"),
                row.get("email"),
            )
        ),
    )


def load_contacts_from_workbook(contacts_path: Path) -> tuple[dict[str, ContactDetails], str] | tuple[None, None]:
    workbook = load_workbook(contacts_path, read_only=True, data_only=True)
    sheet = workbook[workbook.sheetnames[0]]
    headers = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
    normalized_headers = [normalize_space(header).lower() for header in headers]
    header_map = {name: index for index, name in enumerate(normalized_headers)}

    if "pbo_name" not in header_map:
        return None, None

    supported_fields = {
        "contact_name",
        "contact_telephone",
        "contact_email",
        "telephone",
        "cell_phone",
        "email",
    }
    if not any(field in header_map for field in supported_fields):
        return None, None

    grouped_rows: dict[str, list[dict[str, object]]] = {}
    for row in sheet.iter_rows(min_row=2, values_only=True):
        pbo_name = normalize_space(row[header_map["pbo_name"]])
        if not pbo_name:
            continue
        key = normalize_key(pbo_name)
        grouped_rows.setdefault(key, []).append(
            {
                field: row[index]
                for field, index in header_map.items()
                if field in supported_fields
            }
        )

    contacts = {key: merge_contact_rows(rows) for key, rows in grouped_rows.items()}
    return contacts, str(contacts_path.name)


def load_contacts_from_database(database_path: Path) -> tuple[dict[str, ContactDetails], str]:
    connection = sqlite3.connect(database_path)
    connection.row_factory = sqlite3.Row
    cursor = connection.cursor()
    cursor.execute(
        """
        SELECT
            pbo_name,
            contact_name,
            contact_telephone,
            contact_email,
            telephone,
            cell_phone,
            email,
            reporting_period_end,
            id
        FROM pbo_reports
        ORDER BY COALESCE(reporting_period_end, '') DESC, id DESC
        """
    )

    grouped_rows: dict[str, list[dict[str, object]]] = {}
    for record in cursor.fetchall():
        pbo_name = normalize_space(record["pbo_name"])
        if not pbo_name:
            continue
        grouped_rows.setdefault(normalize_key(pbo_name), []).append(dict(record))

    contacts = {key: merge_contact_rows(rows) for key, rows in grouped_rows.items()}
    connection.close()
    return contacts, str(database_path.name)


def choose_contact_source(
    contacts_workbook_path: Path,
    database_path: Path,
) -> tuple[dict[str, ContactDetails], str, bool]:
    workbook_contacts, source_name = load_contacts_from_workbook(contacts_workbook_path)
    if workbook_contacts is not None:
        return workbook_contacts, source_name, False

    database_contacts, source_name = load_contacts_from_database(database_path)
    return database_contacts, source_name, True


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def filter_organizations_by_county(
    organizations: dict[str, ChildOrganization],
    county: str,
) -> dict[str, ChildOrganization]:
    return {
        key: organization
        for key, organization in organizations.items()
        if coverage_includes_county(organization, county)
    }


def build_document(
    output_path: Path,
    organizations: dict[str, ChildOrganization],
    contacts: dict[str, ContactDetails],
    projects_source_name: str,
    contacts_source_name: str,
    used_fallback: bool,
    title_text: str = "Organizations Working With Children",
    selection_note: str | None = None,
) -> tuple[int, int, int]:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(15)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    intro_text = (
        "Organizations were selected from "
        f"{projects_source_name} where the PBO name or aggregated project list contains "
        "'CHILD'."
    )
    if selection_note:
        intro_text = f"{intro_text} {selection_note}"
    intro.add_run(intro_text)

    source_note = document.add_paragraph()
    source_note.paragraph_format.space_after = Pt(8)
    if used_fallback:
        source_note.add_run(
            f"Contact details were not available in the supplied contacts workbook, so the "
            f"latest available contact fields were taken from {contacts_source_name}."
        )
    else:
        source_note.add_run(f"Contact details were taken from {contacts_source_name}.")

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = [
        "Organization",
        "Projects",
        "Coverage",
        "Contact Name",
        "Contact Number",
        "Contact Email",
    ]
    for index, heading in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], heading, bold=True)

    complete = 0
    partial = 0
    missing = 0

    for key, organization in sorted(organizations.items(), key=lambda item: item[1].pbo_name.upper()):
        contact = contacts.get(key, ContactDetails(name="", phone="", email=""))
        present_fields = sum(bool(value) for value in (contact.name, contact.phone, contact.email))
        if present_fields == 3:
            complete += 1
        elif present_fields > 0:
            partial += 1
        else:
            missing += 1

        row = table.add_row().cells
        coverage_parts = [part for part in (organization.scope, organization.counties) if part]
        coverage = " | ".join(coverage_parts) if coverage_parts else "Not specified"

        set_cell_text(row[0], organization.pbo_name)
        set_cell_text(row[1], organization.projects or "Not specified")
        set_cell_text(row[2], coverage)
        set_cell_text(row[3], contact.name or "Not available")
        set_cell_text(row[4], contact.phone or "Not available")
        set_cell_text(row[5], contact.email or "Not available")

    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(8)
    summary.add_run(
        f"Total organizations: {len(organizations)}. "
        f"Complete contact details: {complete}. "
        f"Partial contact details: {partial}. "
        f"No contact details found: {missing}."
    )

    document.save(output_path)
    return complete, partial, missing


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a Word document listing child-focused organizations and contact details."
    )
    parser.add_argument(
        "--projects",
        default="projectIMPLEMENTATIONS.xlsx",
        help="Path to the project implementations workbook.",
    )
    parser.add_argument(
        "--contacts",
        default="pbo_reports (4).xlsx",
        help="Path to the contacts workbook.",
    )
    parser.add_argument(
        "--database",
        default="returnsform14_org_backup.sqlite",
        help="Fallback database path when the contacts workbook does not contain contact fields.",
    )
    parser.add_argument(
        "--output",
        default="children.docx",
        help="Output .docx path.",
    )
    parser.add_argument(
        "--nakuru-output",
        default="children_nakuru.docx",
        help="Output .docx path for child-focused organizations whose coverage includes Nakuru.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    projects_path = Path(args.projects)
    contacts_path = Path(args.contacts)
    database_path = Path(args.database)
    output_path = Path(args.output)
    nakuru_output_path = Path(args.nakuru_output)

    organizations = identify_child_organizations(projects_path)
    contacts, contact_source_name, used_fallback = choose_contact_source(contacts_path, database_path)

    complete, partial, missing = build_document(
        output_path=output_path,
        organizations=organizations,
        contacts=contacts,
        projects_source_name=projects_path.name,
        contacts_source_name=contact_source_name,
        used_fallback=used_fallback,
    )

    nakuru_organizations = filter_organizations_by_county(organizations, "Nakuru")
    nakuru_complete, nakuru_partial, nakuru_missing = build_document(
        output_path=nakuru_output_path,
        organizations=nakuru_organizations,
        contacts=contacts,
        projects_source_name=projects_path.name,
        contacts_source_name=contact_source_name,
        used_fallback=used_fallback,
        title_text="Organizations Working With Children in Nakuru Coverage",
        selection_note="This document further filters the child-focused organizations to those whose coverage includes Nakuru.",
    )

    print(f"Created {output_path}")
    print(f"Organizations: {len(organizations)}")
    print(f"Complete contacts: {complete}")
    print(f"Partial contacts: {partial}")
    print(f"Missing contacts: {missing}")
    print(f"Created {nakuru_output_path}")
    print(f"Nakuru organizations: {len(nakuru_organizations)}")
    print(f"Nakuru complete contacts: {nakuru_complete}")
    print(f"Nakuru partial contacts: {nakuru_partial}")
    print(f"Nakuru missing contacts: {nakuru_missing}")
    if used_fallback:
        print(
            f"{contacts_path.name} does not contain contact columns. "
            f"Used {contact_source_name} instead."
        )


if __name__ == "__main__":
    main()
