from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "pbo_project_implementations__20260402__110536.xlsx"
OUTPUT_DOCX = BASE_DIR / "projectimplementations.docx"
COMPARISON_DOCX = BASE_DIR.parent.parent / "joel comparison.docx"


SECTOR_HARMONIZATION_RULES: list[tuple[str, str]] = [
    (r"\bHIV\s*/?\s*AIDS\b|\bHIVAIDS\b", "HIV/AIDS"),
    (r"\bMICRO\s*-?\s*FINANCE\b|\bMICROFINANCE\b|\bMIRCO-FINANCE\b", "MICROFINANCE"),
    (r"\bENVIRONMENT\b|\bCLIMATE\b|\bCONSERVATION\b", "ENVIRONMENT"),
    (r"\bRELIEF\b|\bDISASTER\b|\bEMERGENCY\b|\bFLOODS?\b|\bDROUGHT\b|\bHUMANITARIAN\b", "RELIEF/DISASTER MANAGEMENT"),
    (r"\bREPRODUCTIVE HEALTH\b|\bPOPULATION AND REPRODUCTIVE HEALTH\b", "REPRODUCTIVE HEALTH"),
    (r"\bYOUTH EMPOWERMENT\b", "YOUTH"),
    (r"\bADVOCACY\b", "ADVOCACY"),
    (r"\bCAPACITY BUILDING\b", "CAPACITY BUILDING"),
]


HISTORICAL_COMPARISON_RULES: list[tuple[str, str]] = [
    (r"\bRESEARCH\b", "Research"),
    (r"\bENERGY\b", "Energy"),
    (r"\bTRANSPORT\b|\bROAD MAINTENANCE\b|\bMOTOR VEHICLE\b", "Transportation"),
    (r"\bROAD SAFETY\b", "Road Safety"),
    (r"\bDRUG\b|\bALCOHOL\b", "Drug and Alcohol Addiction"),
    (r"\bWATER\b|\bSANITATION\b", "Water and Sanitation"),
    (r"\bAGRICULTURE\b|\bPASTORALISM\b", "Agriculture"),
    (r"\bANIMAL WELFARE\b", "Animal Welfare"),
    (r"\bEDUCATION\b|\bCLASSROOM\b|\bLIBRARY\b", "Education"),
    (
        r"\bHEALTH\b|\bHIV\s*/?\s*AIDS\b|\bHIVAIDS\b|\bNUTRITION\b|\bMEDICAL CAMP\b|"
        r"\bREPRODUCTIVE HEALTH\b|\bPOPULATION AND REPRODUCTIVE HEALTH\b",
        "Health",
    ),
    (r"\bENVIRONMENT\b|\bCLIMATE\b|\bCONSERVATION\b", "Environment"),
    (r"\bHOUSING AND SETTLEMENT\b|\bHUMAN SETTLEMENTS?\b", "Housing and Settlement"),
    (r"\bICT\b|\bSCIENCE AND TECHNOLOGY\b|\bINNOVATION\b", "ICT"),
    (r"\bINFORMAL SECTOR\b", "Informal Sector"),
    (r"\bINFORMATION\b|\bPUBLICITY\b|\bMARKETING\b", "Information"),
    (r"\bSPORTS?\b", "Sports"),
    (r"\bDISABIL", "Disability"),
    (r"\bREFUGEES?\b", "Refugees"),
    (r"\bCHILD(REN)?\b", "Children"),
    (r"\bOLD AGE\b", "Old Age Care"),
    (r"\bPEACE BUILDING\b|\bPCVE\b|\bPSVE\b", "Peace Building"),
    (r"\bCULTURE\b|\bPOETRY\b|\bLANGUAGE DEVELOPMENT\b", "Culture"),
    (r"\bGOVERNANCE\b", "Governance"),
    (r"\bHUMAN RIGHTS?\b|\bLAND RIGHTS?\b|\bTAX JUSTICE\b|\bLEGAL\b", "Human Rights"),
    (r"\bGENDER\b|\bWOMEN\b", "Gender"),
    (r"\bYOUTH\b", "Youth"),
    (r"\bMICRO\s*-?\s*FINANCE\b|\bMICROFINANCE\b|\bMIRCO-FINANCE\b", "Micro-Finance"),
    (
        r"\bCAPACITY BUILDING\b|\bCOMMUNITY CAPACITY BUILDINGS?\b|\bORGANI[ZS]ATIONAL DEVELOPMENT\b|"
        r"\bPROJECT MONITORING.*CAPACITY BUILDING\b|\bMONITORING AND EVALUATION\b|\bMETRICS AND EVALUATION\b",
        "Capacity Building",
    ),
    (r"\bADVOCACY\b|\bEMPOWERMENT\b|\bLEADERSHIP TRAINING\b|\bCOMMUNITY EMPOWERMENT\b", "Advocacy and Empowerment"),
    (r"\bRELIGI|\bCHURCH\b|\bBIBLE\b|\bPASTORS?\b|\bMOSQUES?\b|\bRAMATHAN\b|\bQURBANI\b|\bSPIRITUAL\b", "Religion"),
    (
        r"\bRELIEF\b|\bDISASTER\b|\bEMERGENCY\b|\bHUMANITARIAN\b|\bFOOD SECURITY\b|"
        r"\bPROTECTION AND ASSISTANCE IN EMERGENCY\b|\bPROTECTION AND SOCIAL COHESION\b|"
        r"\bPROTECTION FROM ALL FORM OF VIOLENCE\b",
        "Relief/Disaster Management",
    ),
    (
        r"\bWELFARE\b|\bVULNERABLE\b|\bWIDOWS?\b|\bCHARITABLE CONTRIBUTION\b|\bCARE REFORMS\b|"
        r"\bSOCIAL PROTECTION\b|\bADEQUATE STANDARD OF LIVING\b|\bSOCIAL SERVICES\b",
        "Welfare",
    ),
    (
        r"\bDEVELOPMENT\b|\bLIVELIHOOD\b|\bPOVERTY ERADICATION\b|"
        r"\bPOVERTY ALLEVIATION\b|\bEXTREME POVERTY\b|\bECONOMIC RECOVERY\b|\bSUSTAINABLE DEVELOPMENT\b|"
        r"\bHUMAN CAPITAL DEVELOPMENT\b|\bCOMMUNITY SUPPORT\b|\bECONOMIC INCLUSION\b",
        "Development",
    ),
]


MULTI_SECTORAL_RULES = (
    r"\bINTERGRATED\b|\bINTEGRATED\b|\bQUALITY SERVICES\b|\bQUALITY PROGRAMMING\b|\bNOT SPECIFIED\b|"
    r"\bGENERAL DIRECT EXPENDITURE\b|\bGENERAL EXPENSES\b|\bADMINISTRATION\b|\bADMINISTRATIVE\b|\bAUDIT FEES\b|"
    r"\bFEES\b|\bHQ UNRESTRICTED\b|\bPROJECT VEHICLES\b|\bFIELD OFFICE\b|\bSUGRANTING TO PARTNERS\b|"
    r"\bRENT\b|\bINFRASTRUCTURE\b|\bCONSTRUCTION\b|\bREPAIRS\b|\bPROJECT ADMINISTRATION\b|\bPETTY CASH\b"
)


ORG_DISPLAY_OVERRIDES = {
    "ASSOCIATION OF CHRISTIAN RESOURCE ORGANIZATIONS SERVING SUDAN ACROSS": "ASSOCIATION OF CHRISTIAN RESOURCE ORGANIZATIONS SERVING SUDAN (ACROSS)",
    "CENTER FOR INTERNATIONAL HEALTH EDUCATION AND BIOSECURITY CIHEB KENYA": "CENTER FOR INTERNATIONAL HEALTH, EDUCATION AND BIOSECURITY (CIHEB)-KENYA",
    "SAMARITAN S PURSE": "SAMARITAN'S PURSE",
    "CATHOLIC RELIEF SERVICES USCC": "CATHOLIC RELIEF SERVICES - USCC",
    "WORLD WIDE FUND FOR NATURE KENYA WWF KENYA": "WORLD WIDE FUND FOR NATURE KENYA (WWF-KENYA)",
    "UNITED WOMEN EMPOWERMENT PROGRAMME AFRICA": "UNITED WOMEN EMPOWERMENT PROGRAMME-AFRICA",
}


COUNTY_NORMALIZATION_MAP = {
    "BARINGO": "BARINGO",
    "BOMET": "BOMET",
    "BUNGOMA": "BUNGOMA",
    "BUSIA": "BUSIA",
    "ELGEYO MARAKWET": "ELGEYO MARAKWET",
    "EMBU": "EMBU",
    "GARISSA": "GARISSA",
    "HOMA BAY": "HOMA BAY",
    "ISIOLO": "ISIOLO",
    "KAJIADO": "KAJIADO",
    "KAKAMEGA": "KAKAMEGA",
    "KERICHO": "KERICHO",
    "KIAMBU": "KIAMBU",
    "KILIFI": "KILIFI",
    "KIRINYAGA": "KIRINYAGA",
    "KISII": "KISII",
    "KISUMU": "KISUMU",
    "KITUI": "KITUI",
    "KWALE": "KWALE",
    "LAIKIPIA": "LAIKIPIA",
    "LAMU": "LAMU",
    "MACHAKOS": "MACHAKOS",
    "MAKUENI": "MAKUENI",
    "MANDERA": "MANDERA",
    "MARSABIT": "MARSABIT",
    "MERU": "MERU",
    "MIGORI": "MIGORI",
    "MOMBASA": "MOMBASA",
    "MURANGA": "MURANG'A",
    "MURANG'A": "MURANG'A",
    "NAIROBI": "NAIROBI CITY",
    "NAIROBI CITY": "NAIROBI CITY",
    "NAKURU": "NAKURU",
    "NANDI": "NANDI",
    "NAROK": "NAROK",
    "NYAMIRA": "NYAMIRA",
    "NYANDARUA": "NYANDARUA",
    "NYERI": "NYERI",
    "SAMBURU": "SAMBURU",
    "SIAYA": "SIAYA",
    "TAITA TAVETA": "TAITA TAVETA",
    "TANA RIVER": "TANA RIVER",
    "THARAKA NITHI": "THARAKA NITHI",
    "TRANS NZOIA": "TRANS NZOIA",
    "TRANS-NZOIA": "TRANS NZOIA",
    "TURKANA": "TURKANA",
    "UASIN GISHU": "UASIN GISHU",
    "VIHIGA": "VIHIGA",
    "WAJIR": "WAJIR",
    "WEST POKOT": "WEST POKOT",
    "ALL COUNTIES": "ALL COUNTIES",
}


def normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def clean_sector(value) -> str:
    raw = normalize_whitespace("" if pd.isna(value) else value).upper()
    if not raw:
        return "UNSPECIFIED"
    for pattern, replacement in SECTOR_HARMONIZATION_RULES:
        if re.search(pattern, raw):
            return replacement
    return raw


def map_sector_to_historical(value) -> str:
    raw = normalize_whitespace("" if pd.isna(value) else value).upper()
    if not raw:
        return "Multi - Sectoral"
    for pattern, replacement in HISTORICAL_COMPARISON_RULES:
        if re.search(pattern, raw):
            return replacement
    if re.search(MULTI_SECTORAL_RULES, raw):
        return "Multi - Sectoral"
    return "Multi - Sectoral"


def canonical_org_key(name, normalized_name) -> str:
    base = normalized_name if isinstance(normalized_name, str) and normalized_name.strip() else name
    base = normalize_whitespace(base).upper()
    base = base.replace("&", " AND ")
    base = re.sub(r"[^A-Z0-9\s]", " ", base)
    base = re.sub(r"\s+", " ", base).strip()

    if "ASSOCIATION OF CHRISTIAN RESOURCE ORGANIZATIONS" in base and "ACROSS" in base:
        return "ASSOCIATION OF CHRISTIAN RESOURCE ORGANIZATIONS SERVING SUDAN ACROSS"
    return base


def org_display_name(group: pd.DataFrame, key: str) -> str:
    if key in ORG_DISPLAY_OVERRIDES:
        return ORG_DISPLAY_OVERRIDES[key]
    names = group["pbo_name"].dropna().astype(str).map(normalize_whitespace)
    if names.empty:
        return key.title()
    counts = names.value_counts()
    return str(counts.index[0])


def money(value: float) -> str:
    return f"KES {value:,.2f}"


def pct(value: float) -> str:
    return f"{value:.2f}%"


def fmt_change(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "-"
    return f"{value:.2f}%"


def parse_amount(value) -> float:
    text = normalize_whitespace("" if pd.isna(value) else value)
    if not text or text == "-":
        return 0.0
    return float(text.replace(",", ""))


def change_pct(current: float, previous: float) -> float | None:
    if previous == 0:
        return None
    return (current - previous) / previous * 100


def load_historical_sector_table() -> pd.DataFrame:
    if not COMPARISON_DOCX.exists():
        raise FileNotFoundError(f"Comparison document not found: {COMPARISON_DOCX}")

    document = Document(COMPARISON_DOCX)
    if not document.tables:
        raise ValueError(f"No comparison table found in {COMPARISON_DOCX}")

    rows = []
    for row in document.tables[0].rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 3:
            continue
        rows.append(
            {
                "Sector": cells[0],
                "FY 2023/2024": parse_amount(cells[1]),
                "FY 2022/23": parse_amount(cells[2]),
            }
        )
    return pd.DataFrame(rows)


def split_counties(value) -> list[str]:
    if pd.isna(value):
        return []
    text = normalize_whitespace(str(value)).upper()
    if not text or text == "NAN":
        return []

    counties = []
    for part in re.split(r"\s*,\s*|\s*;\s*|\s*/\s*|\s+AND\s+", text):
        part = normalize_whitespace(part)
        if not part:
            continue
        counties.append(COUNTY_NORMALIZATION_MAP.get(part, part))
    return sorted(set(counties))


def build_county_analysis_table(df: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, int]]:
    county_counts: dict[str, int] = {}
    all_counties_rows = 0
    unspecified_rows = 0
    specific_county_rows = 0

    for value in df["counties"]:
        counties = split_counties(value)
        if not counties:
            unspecified_rows += 1
            continue

        if "ALL COUNTIES" in counties:
            all_counties_rows += 1
            counties = [county for county in counties if county != "ALL COUNTIES"]

        if not counties:
            continue

        specific_county_rows += 1
        for county in counties:
            county_counts[county] = county_counts.get(county, 0) + 1

    county_table = (
        pd.DataFrame(
            [
                {"county": county, "projects_implemented": count}
                for county, count in county_counts.items()
            ]
        )
        .sort_values(["projects_implemented", "county"], ascending=[False, True])
        .reset_index(drop=True)
    )
    county_table["projects_implemented"] = county_table["projects_implemented"].astype(int)
    county_table["share_of_project_rows"] = (
        county_table["projects_implemented"] / len(df) * 100
    ).round(2)

    summary = {
        "rows_marked_all_counties": all_counties_rows,
        "rows_with_no_county": unspecified_rows,
        "rows_with_specific_counties": specific_county_rows,
    }
    return county_table, summary


def build_top_spend_table(df: pd.DataFrame, amount_col: str, total_amount: float) -> pd.DataFrame:
    work = df.copy()
    work["_amount"] = pd.to_numeric(work[amount_col], errors="coerce").fillna(0.0)
    work["_org_key"] = [
        canonical_org_key(name, normalized)
        for name, normalized in zip(work["pbo_name"], work["pbo_name_normalized"])
    ]

    grouped = []
    for key, group in work.groupby("_org_key", dropna=False):
        total = float(group["_amount"].sum())
        rows = int((group["_amount"] > 0).sum())
        if total <= 0:
            continue
        grouped.append(
            {
                "organization": org_display_name(group, key),
                "frequency": rows,
                "total_amount": total,
                "percentage": (total / total_amount * 100) if total_amount else 0.0,
            }
        )

    result = pd.DataFrame(grouped).sort_values(
        ["total_amount", "frequency", "organization"],
        ascending=[False, False, True],
    ).head(20).reset_index(drop=True)
    result.index = result.index + 1
    return result


def build_sector_table(df: pd.DataFrame) -> tuple[pd.DataFrame, int]:
    work = df.copy()
    work["_sector"] = work["sector"].map(clean_sector)
    work["_sector_amount"] = pd.to_numeric(work["amount_spent_kenya"], errors="coerce").fillna(0.0)
    grand_total = float(work["_sector_amount"].sum())
    total_rows = len(work)
    table = (
        work.groupby("_sector", dropna=False)
        .agg(
            total_amount=("_sector_amount", "sum"),
        )
        .reset_index()
        .rename(columns={"_sector": "sector"})
        .sort_values(["total_amount", "sector"], ascending=[False, True])
        .reset_index(drop=True)
    )
    table["percentage"] = (
        table["total_amount"] / grand_total * 100 if grand_total else 0.0
    ).round(2)
    return table, total_rows


def build_sector_spend_comparison(df: pd.DataFrame, historical_table: pd.DataFrame) -> pd.DataFrame:
    work = df.copy()
    work["_fy_24_25"] = pd.to_numeric(work["amount_spent_kenya"], errors="coerce").fillna(0.0)
    work["_historical_sector"] = work["sector"].map(map_sector_to_historical)
    current_totals = (
        work.groupby("_historical_sector", dropna=False)["_fy_24_25"]
        .sum()
        .to_dict()
    )
    kenya_total = float(work["_fy_24_25"].sum())

    rows = []
    for _, historical_row in historical_table.iterrows():
        sector = str(historical_row["Sector"])
        previous_year = float(historical_row["FY 2023/2024"])
        earlier_year = float(historical_row["FY 2022/23"])
        if sector.upper() == "TOTAL":
            current_year = kenya_total
        else:
            current_year = float(current_totals.get(sector, 0.0))

        rows.append(
            {
                "Sector": sector,
                "FY 24/25": current_year,
                "FY 2023/2024": previous_year,
                "FY 2022/23": earlier_year,
                "FY 24/25 vs FY 2023/2024": fmt_change(change_pct(current_year, previous_year)),
                "FY 2023/2024 vs FY 2022/23": fmt_change(change_pct(previous_year, earlier_year)),
            }
        )

    return pd.DataFrame(rows)


def add_title(document: Document, text: str):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(document: Document, text: str):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)


def add_bullet(document: Document, text: str):
    document.add_paragraph(text, style="List Bullet")


def add_dataframe_table(document: Document, df: pd.DataFrame, column_formats: dict[str, str] | None = None):
    if df.empty:
        document.add_paragraph("No data available.")
        return
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_cells[idx].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            value = row[col]
            fmt = (column_formats or {}).get(col)
            if fmt == "money":
                cells[idx].text = money(float(value))
            elif fmt == "pct":
                cells[idx].text = pct(float(value))
            else:
                cells[idx].text = str(value)


def main():
    df = pd.read_excel(SOURCE_PATH, sheet_name="Matched Rows")
    historical_sector_table = load_historical_sector_table()

    sector_table, total_rows = build_sector_table(df)
    county_table, county_summary = build_county_analysis_table(df)
    kenya_total = float(pd.to_numeric(df["amount_spent_kenya"], errors="coerce").fillna(0.0).sum())
    sector_comparison_table = build_sector_spend_comparison(df, historical_sector_table)

    kenya_top = build_top_spend_table(df, "amount_spent_kenya", kenya_total)
    kenya_summary_table = pd.DataFrame(
        [
            {
                "amount_type": "Total amount spent in Kenya",
                "total_amount": kenya_total,
            }
        ]
    )

    full_sector_table = sector_table.copy()
    previous_total_row = sector_comparison_table[sector_comparison_table["Sector"].str.upper() == "TOTAL"].iloc[0]
    current_comparison_leader = (
        sector_comparison_table[sector_comparison_table["Sector"].str.upper() != "TOTAL"]
        .sort_values("FY 24/25", ascending=False)
        .iloc[0]
    )
    leading_county = county_table.iloc[0]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(document, "Project Implementations Analysis Report")
    add_subtitle(document, f"Source workbook: {SOURCE_PATH.name}")
    add_subtitle(document, f"Historical comparison source: {COMPARISON_DOCX.name}")
    add_subtitle(document, f"Matched project rows analyzed: {total_rows:,}")

    document.add_heading("Cleaning Approach", level=1)
    add_bullet(document, "Sector values were trimmed, uppercased, and whitespace-normalized.")
    add_bullet(document, "Close sector variants were harmonized for the sector table, including HIV/AIDS, MICROFINANCE, ENVIRONMENT, RELIEF/DISASTER MANAGEMENT, REPRODUCTIVE HEALTH, and YOUTH.")
    add_bullet(document, "FY 24/25 sector amounts in the sector table were calculated from amount_spent_kenya only.")
    add_bullet(document, "The UNSPECIFIED sector represents rows where the source sector field was blank.")
    add_bullet(document, "For the historical comparison table, current sector labels were remapped into the older comparison taxonomy used in joel comparison.docx.")
    add_bullet(document, "Blank, administrative, and cross-cutting current sector labels were rolled into Multi - Sectoral in the historical comparison table so the FY 24/25 Kenya totals reconcile to the current dataset.")
    add_bullet(document, "County analysis uses the populated counties field because the single county column is empty in this workbook.")
    add_bullet(document, "Organization rankings were grouped using pbo_name_normalized where available, then punctuation and spacing were normalized to reduce duplicates.")

    document.add_heading("Sector Frequency Analysis", level=1)
    document.add_paragraph(
        "The table below lists all cleaned sector values by FY 24/25 amount spent using amount_spent_kenya only. The percentage column shows each sector's share of the total FY 24/25 Kenya spend across all sectors."
    )
    sector_frequency_display = full_sector_table.rename(
        columns={
            "sector": "Sector",
            "total_amount": "FY 24/25 total amount",
            "percentage": "Percentage of total amount",
        }
    )[
        ["Sector", "FY 24/25 total amount", "Percentage of total amount"]
    ]
    add_dataframe_table(
        document,
        sector_frequency_display,
        column_formats={
            "FY 24/25 total amount": "money",
            "Percentage of total amount": "pct",
        },
    )

    document.add_heading("Sector Spending Comparison with Previous Years", level=1)
    document.add_paragraph(
        "The table below compares current FY 24/25 sector spending using amount_spent_kenya only with the previous-year sector table from joel comparison.docx. FY 24/25 is placed immediately after Sector for direct reading across years."
    )
    add_dataframe_table(
        document,
        sector_comparison_table,
        column_formats={
            "FY 24/25": "money",
            "FY 2023/2024": "money",
            "FY 2022/23": "money",
        },
    )

    document.add_heading("County Analysis", level=1)
    document.add_paragraph(
        "The table below counts how many project rows mention each county. Because the source workbook uses a multi-value counties field, one project row can contribute to more than one county, so county totals are not mutually exclusive."
    )
    document.add_paragraph(
        f"Rows with specific counties: {county_summary['rows_with_specific_counties']:,}. "
        f"Rows marked ALL COUNTIES only: {county_summary['rows_marked_all_counties']:,}. "
        f"Rows with no county listed: {county_summary['rows_with_no_county']:,}."
    )
    add_dataframe_table(
        document,
        county_table.rename(
            columns={
                "county": "County",
                "projects_implemented": "Projects implemented",
                "share_of_project_rows": "Share of project rows",
            }
        ),
        column_formats={"Share of project rows": "pct"},
    )

    document.add_heading("Top 20 Organizations by Amount Spent in Kenya", level=1)
    document.add_paragraph(
        "Frequency counts here refer to the number of project rows for each organization with a positive amount_spent_kenya value."
    )
    add_dataframe_table(
        document,
        kenya_top.reset_index().rename(
            columns={
                "index": "Rank",
                "organization": "Organization",
                "frequency": "Frequency",
                "total_amount": "Total amount_spent_kenya",
                "percentage": "Share of Kenya total",
            }
        ),
        column_formats={
            "Total amount_spent_kenya": "money",
            "Share of Kenya total": "pct",
        },
    )

    document.add_heading("Kenya Amount Summary", level=1)
    document.add_paragraph(
        "The summary below reports the total FY 24/25 amount_spent_kenya used across the Kenya-only amount sections in this report."
    )
    add_dataframe_table(
        document,
        kenya_summary_table.rename(
            columns={
                "amount_type": "Amount type",
                "total_amount": "Total amount",
            }
        ),
        column_formats={
            "Total amount": "money",
        },
    )

    document.add_page_break()
    document.add_heading("Headline Findings", level=1)
    document.add_paragraph(
        f"{full_sector_table.iloc[0]['sector'].title()} leads the sector table by FY 24/25 spend "
        f"with {money(float(full_sector_table.iloc[0]['total_amount']))} "
        f"({pct(float(full_sector_table.iloc[0]['percentage']))} of total Kenya sector spend)."
    )
    document.add_paragraph(
        f"The largest FY 24/25 sector in the historical comparison view is {current_comparison_leader['Sector']} "
        f"at {money(float(current_comparison_leader['FY 24/25']))}."
    )
    document.add_paragraph(
        f"Nairobi City is the most frequently mentioned county, appearing in {int(leading_county['projects_implemented']):,} project rows "
        f"({pct(float(leading_county['share_of_project_rows']))})."
    )
    document.add_paragraph(
        f"The total FY 24/25 Kenya spend is {money(kenya_total)} compared with {money(float(previous_total_row['FY 2023/2024']))} in FY 2023/2024, "
        f"a change of {previous_total_row['FY 24/25 vs FY 2023/2024']}."
    )
    document.add_paragraph(
        f"The total amount_spent_kenya captured in this report is {money(kenya_total)}."
    )
    document.add_paragraph(
        f"The leading organization by amount_spent_kenya is {kenya_top.iloc[0]['organization']} "
        f"with {money(float(kenya_top.iloc[0]['total_amount']))}."
    )

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
