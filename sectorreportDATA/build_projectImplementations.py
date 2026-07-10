from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "edited sector spending 24.25.xlsx"
OUTPUT_DOCX = BASE_DIR / "projectimplementations.docx"


def normalize_whitespace(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def normalize_sector_text(value: object) -> str:
    if pd.isna(value):
        return ""

    text = str(value).upper()
    text = text.replace("&", " AND ")
    text = text.replace("/", " ")
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    return normalize_whitespace(text)


def money(value: float) -> str:
    return f"KES {value:,.2f}"


def pct(value: float) -> str:
    return f"{value:.2f}%"


def normalize_scope(value: object) -> str:
    text = normalize_whitespace(value).lower()
    if text in {"nngo", "national"}:
        return "National"
    if text in {"ingo", "international"}:
        return "International"
    return "Unspecified"


def load_source_data() -> pd.DataFrame:
    df = pd.read_excel(SOURCE_PATH, sheet_name="Sheet1").copy()
    df["pbo_name"] = df["pbo_name"].fillna("Unknown organization").astype(str).map(normalize_whitespace)
    df["raw_sector_display"] = (
        df["sector"]
        .fillna("Blank / missing sector")
        .astype(str)
        .map(normalize_whitespace)
        .replace("", "Blank / missing sector")
    )
    df["normalized_sector"] = df["sector"].map(normalize_sector_text).replace("", "BLANK MISSING SECTOR")
    df["sector_label"] = df["raw_sector_display"]
    df["scope_label"] = df.get("scope", pd.Series(index=df.index, dtype="object")).map(normalize_scope)
    df["amount_spent_kenya"] = pd.to_numeric(df["amount_spent_kenya"], errors="coerce").fillna(0.0)
    return df


def build_sector_spend_table(df: pd.DataFrame) -> pd.DataFrame:
    grand_total = float(df["amount_spent_kenya"].sum())
    table = (
        df.groupby("raw_sector_display", dropna=False)["amount_spent_kenya"]
        .sum()
        .reset_index()
        .rename(
            columns={
                "raw_sector_display": "Sector",
                "amount_spent_kenya": "FY 24/25 Kenya amount",
            }
        )
    )
    table["Share of FY 24/25 Kenya amount"] = (
        table["FY 24/25 Kenya amount"] / grand_total * 100 if grand_total else 0.0
    ).round(2)
    table = table.sort_values(
        ["FY 24/25 Kenya amount", "Sector"],
        ascending=[False, True],
    ).reset_index(drop=True)

    total_row = pd.DataFrame(
        [
            {
                "Sector": "TOTAL",
                "FY 24/25 Kenya amount": grand_total,
                "Share of FY 24/25 Kenya amount": 100.0 if grand_total else 0.0,
            }
        ]
    )
    return pd.concat([table, total_row], ignore_index=True)


def build_raw_mapping_table(df: pd.DataFrame, top_n: int = 100) -> pd.DataFrame:
    return (
        df.groupby(["raw_sector_display", "normalized_sector", "sector_label"], dropna=False)
        .agg(
            fy_24_25_kenya_amount=("amount_spent_kenya", "sum"),
            rows=("sector_label", "size"),
        )
        .reset_index()
        .sort_values(
            ["fy_24_25_kenya_amount", "rows", "raw_sector_display"],
            ascending=[False, False, True],
        )
        .head(top_n)
        .rename(
            columns={
                "raw_sector_display": "Source sector label",
                "normalized_sector": "Normalized sector label",
                "sector_label": "Assigned sector",
                "fy_24_25_kenya_amount": "FY 24/25 Kenya amount",
                "rows": "Rows",
            }
        )
        .reset_index(drop=True)
    )


def build_top_ngo_table(df: pd.DataFrame, scope_label: str, top_n: int = 50) -> pd.DataFrame:
    scope_df = df[df["scope_label"] == scope_label].copy()
    total_amount = float(scope_df["amount_spent_kenya"].sum())
    if scope_df.empty:
        return pd.DataFrame(
            columns=[
                "Organization",
                "Scope",
                "FY 24/25 Kenya amount",
                "Share of scope total",
                "Primary sector",
            ]
        )

    grouped = (
        scope_df.groupby("pbo_name", dropna=False)
        .agg(
            **{
                "FY 24/25 Kenya amount": ("amount_spent_kenya", "sum"),
                "Scope": ("scope_label", "first"),
            }
        )
        .reset_index()
        .rename(columns={"pbo_name": "Organization"})
        .sort_values(["FY 24/25 Kenya amount", "Organization"], ascending=[False, True])
        .head(top_n)
        .reset_index(drop=True)
    )

    primary_sector = (
        scope_df.groupby(["pbo_name", "sector_label"], dropna=False)["amount_spent_kenya"]
        .sum()
        .reset_index()
        .sort_values(["pbo_name", "amount_spent_kenya", "sector_label"], ascending=[True, False, True])
        .drop_duplicates(subset=["pbo_name"])
        .rename(columns={"pbo_name": "Organization", "sector_label": "Primary sector"})
        [["Organization", "Primary sector"]]
    )

    grouped = grouped.merge(primary_sector, on="Organization", how="left")
    grouped["Share of scope total"] = (
        grouped["FY 24/25 Kenya amount"] / total_amount * 100 if total_amount else 0.0
    ).round(2)

    return grouped[
        [
            "Organization",
            "Scope",
            "FY 24/25 Kenya amount",
            "Share of scope total",
            "Primary sector",
        ]
    ]


def add_title(document: Document, text: str) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(document: Document, text: str) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_dataframe_table(
    document: Document, df: pd.DataFrame, column_formats: dict[str, str] | None = None
) -> None:
    if df.empty:
        document.add_paragraph("No data available.")
        return

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            value = row[col]
            fmt = (column_formats or {}).get(col)
            if fmt == "money":
                cells[idx].text = money(float(value))
            elif fmt == "pct":
                cells[idx].text = pct(float(value))
            else:
                cells[idx].text = str(value)


def build_report(
    df: pd.DataFrame,
    sector_table: pd.DataFrame,
    mapping_table: pd.DataFrame,
    national_top_table: pd.DataFrame,
    international_top_table: pd.DataFrame,
) -> Document:
    grand_total = float(df["amount_spent_kenya"].sum())
    positive_sector_table = sector_table[sector_table["Sector"] != "TOTAL"].copy()
    top_sector = positive_sector_table.iloc[0]
    unique_raw_labels = int(df["raw_sector_display"].nunique())
    national_total = float(df.loc[df["scope_label"] == "National", "amount_spent_kenya"].sum())
    international_total = float(df.loc[df["scope_label"] == "International", "amount_spent_kenya"].sum())

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(document, "FY 24/25 Kenya Sector Spend Analysis")
    add_subtitle(document, f"Source workbook: {SOURCE_PATH.name}")
    add_subtitle(document, f"Matched project rows analyzed: {len(df):,}")

    document.add_heading("Methodology", level=1)
    add_bullet(document, "Sector analysis uses source sector labels directly from the workbook without regrouping into predefined sector bands.")
    add_bullet(document, "A normalized sector text field is retained for traceability and quality review, but totals are computed using the original source sector labels.")
    add_bullet(document, "This report analyzes amount_spent_kenya only.")
    add_bullet(document, "Scope values in the source workbook were normalized so NNGO and national are reported as National, while INGO and international are reported as International.")

    document.add_heading("FY 24/25 Kenya Sector Utilization", level=1)
    document.add_paragraph(
        "The table below shows sector-level FY 24/25 Kenya spending based on source sector labels from the workbook."
    )
    add_dataframe_table(
        document,
        sector_table,
        column_formats={
            "FY 24/25 Kenya amount": "money",
            "Share of FY 24/25 Kenya amount": "pct",
        },
    )

    document.add_heading("Headline Findings", level=1)
    document.add_paragraph(
        f"The largest FY 24/25 Kenya sector is {top_sector['Sector']} at "
        f"{money(float(top_sector['FY 24/25 Kenya amount']))}, representing "
        f"{pct(float(top_sector['Share of FY 24/25 Kenya amount']))} of total FY 24/25 Kenya spend."
    )
    document.add_paragraph(
        f"The total FY 24/25 Kenya amount captured in the workbook is {money(grand_total)}."
    )
    document.add_paragraph(
        f"{unique_raw_labels:,} distinct source sector labels were identified in the workbook."
    )
    document.add_paragraph(
        f"National NGOs account for {money(national_total)} and international NGOs account for {money(international_total)} in FY 24/25 Kenya spending based on the workbook scope field."
    )

    document.add_heading("Top 50 National NGO Utilization", level=1)
    document.add_paragraph(
        "The table below ranks the top 50 national NGOs by FY 24/25 Kenya utilization of funds using the scope values in the edited workbook."
    )
    add_dataframe_table(
        document,
        national_top_table,
        column_formats={
            "FY 24/25 Kenya amount": "money",
            "Share of scope total": "pct",
        },
    )

    document.add_heading("Top 50 International NGO Utilization", level=1)
    document.add_paragraph(
        "The table below ranks the top 50 international NGOs by FY 24/25 Kenya utilization of funds using the scope values in the edited workbook."
    )
    add_dataframe_table(
        document,
        international_top_table,
        column_formats={
            "FY 24/25 Kenya amount": "money",
            "Share of scope total": "pct",
        },
    )

    document.add_page_break()
    document.add_heading("Top Normalized Sector Mappings", level=1)
    document.add_paragraph(
        "This appendix shows the highest-value source sector labels, their normalized form, and the final assigned sector."
    )
    add_dataframe_table(
        document,
        mapping_table,
        column_formats={"FY 24/25 Kenya amount": "money"},
    )

    return document


def main() -> None:
    df = load_source_data()
    sector_table = build_sector_spend_table(df)
    mapping_table = build_raw_mapping_table(df)
    national_top_table = build_top_ngo_table(df, "National")
    international_top_table = build_top_ngo_table(df, "International")
    document = build_report(
        df,
        sector_table,
        mapping_table,
        national_top_table,
        international_top_table,
    )
    document.save(OUTPUT_DOCX)

    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
