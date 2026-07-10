from __future__ import annotations

import re
import subprocess
from pathlib import Path

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "collaboration.xlsx"
OUTPUT_DOCX = BASE_DIR / "collaboration.docx"
REPORTS_DIR = BASE_DIR / "reports"
ASSETS_DIR = REPORTS_DIR / "collaboration_assets"
OUTPUT_HTML = REPORTS_DIR / "collaboration_report.html"

COLLABORATION_COLUMNS = [
    "info_exchange",
    "tech_support_to_partner",
    "tech_support_from_partner",
    "funding_to_partner",
    "funding_from_partner",
    "equipment_to_partner",
    "equipment_from_partner",
]

COLLABORATION_LABELS = {
    "info_exchange": "Information exchange",
    "tech_support_to_partner": "Technical support to partner",
    "tech_support_from_partner": "Technical support from partner",
    "funding_to_partner": "Funding to partner",
    "funding_from_partner": "Funding from partner",
    "equipment_to_partner": "Equipment to partner",
    "equipment_from_partner": "Equipment from partner",
}

PARTNER_GROUP_RULES: list[tuple[str, str]] = [
    ("PBOs", r"^PBOS?$|CARIBOU DIGITAL|NGO"),
    ("CBOs", r"^CBOS?$"),
    ("FBOs / Religious Institutions", r"^FBOS?$|RELIGIOUS|CHURCH|CHURCH INSTITUTIONS|MOSQUE|CATHOLIC|DIOCESE|DIOCESAN|FAITH-BASED"),
    ("Government Agencies", r"GOVERNMENT|MINISTRY| COUNTY|NATIONAL TREASURY|STATE DEPARTMENT|PUBLIC HEALTH|PUBLIC INSTITUTIONS|PUBLIC SCHOOLS"),
    ("Academic Institutions", r"ACADEMIC INSTITUTIONS|SCHOOLS|RESEARCH STUDENTS| KPSA"),
    ("Health Institutions", r"HEALTH INSTITUTIONS|HOSPITAL|LEVEL 4|LEVEL 5| HEALTH| HIV"),
    ("Research Institutions", r"RESEARCH INSTITUTIONS| RESEARCH"),
    ("Donor Agencies", r"DONOR AGENCIES|SEGAL FAMILY FOUNDATION|GIZ"),
    ("Media", r"^MEDIA?$|MEDIA HOUSES"),
    (
        "Corporate",
        r"CORPORATE|CORPORATES|LIMITED COMPANY|SAFARICOM|GOOGLE|META|FACEBOOK|UBER|CARIBOU DIGITAL|FINANCIAL INSTITUTION|FINANCIAL INSTITUTIONS"
        r"CONSULTANCY|FINANCIAL INSTITUTION|MASTERCARD|PRIVATE SOOCIAL ENTERPRISES|SEGAL FAMILY FOUNDATION|AGRO DEALERS",
    ),
    (
        "Others",
        r"SELF HELP GROUP|ORPHANS AND WOMEN GROUP|PARENTS OF CHILDREN WITH ASD|INDIVUDUAL| INDIVIDUAL DONATIONS| KILIMANJARO BLIN TRUST"
        r"WELL WISHERS|PRIVATE INDIVIDUALS|SELF HELP| NOT INDICATED|OTHER| |INDIVIDUAL DONATION| PRIVATE| STTG",
    ),
]

PLOTLY_COLORS = [
    "#0B3954",
    "#087E8B",
    "#BFD7EA",
    "#FF5A5F",
    "#C81D25",
    "#F4D35E",
    "#2D6A4F",
]


def normalize_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def pbo_key(value) -> str:
    return normalize_text(value).upper()


def is_flagged(series: pd.Series) -> pd.Series:
    return series.notna() & series.astype(str).str.strip().ne("")


def map_partner_group(raw_value) -> str:
    raw = normalize_text(raw_value).upper()
    if not raw:
        return "Missing partner type"
    for label, pattern in PARTNER_GROUP_RULES:
        if re.search(pattern, raw):
            return label
    return "Missing Partners Type"


def pct(value: float) -> str:
    return f"{value:.2f}%"


def add_title(document: Document, text: str):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(document: Document, text: str):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)


def add_bullet(document: Document, text: str):
    document.add_paragraph(text, style="List Bullet")


def add_dataframe_table(document: Document, df: pd.DataFrame):
    if df.empty:
        document.add_paragraph("No data available.")
        return

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            cells[idx].text = str(row[col])


def add_figure(document: Document, image_path: Path, caption: str, width: float = 6.7):
    if not image_path.exists() or image_path.stat().st_size == 0:
        document.add_paragraph(f"[Chart unavailable: {caption}]")
        return
    document.add_picture(str(image_path), width=Inches(width))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def prepare_data() -> dict[str, pd.DataFrame | int | float]:
    df = pd.read_excel(SOURCE_PATH, sheet_name=0)
    df["pbo_key"] = df["pbo_name"].map(pbo_key)
    df["partner_type_raw"] = df["partner_type"].fillna("").astype(str).map(normalize_text)
    df["partner_group"] = df["partner_type_raw"].map(map_partner_group)

    flag_frame = pd.DataFrame({col: is_flagged(df[col]) for col in COLLABORATION_COLUMNS})
    df["has_any_collaboration"] = flag_frame.any(axis=1)

    total_pbos = int(df["pbo_key"].nunique())
    collaborating_pbos = int(df.loc[df["has_any_collaboration"], "pbo_key"].nunique())
    no_recorded_collaboration = total_pbos - collaborating_pbos

    coverage_table = pd.DataFrame(
        [
            {
                "Status": "Collaborated with at least one partner",
                "PBOs": collaborating_pbos,
                "Percentage of all PBOs": pct(collaborating_pbos / total_pbos * 100 if total_pbos else 0.0),
            },
            {
                "Status": "No collaboration type recorded",
                "PBOs": no_recorded_collaboration,
                "Percentage of all PBOs": pct(no_recorded_collaboration / total_pbos * 100 if total_pbos else 0.0),
            },
            {
                "Status": "Total PBOs in collaboration workbook",
                "PBOs": total_pbos,
                "Percentage of all PBOs": pct(100.0 if total_pbos else 0.0),
            },
        ]
    )

    collaborating = df.loc[df["has_any_collaboration"]].copy()

    partner_group_table = (
        collaborating.groupby("partner_group")["pbo_key"]
        .nunique()
        .sort_values(ascending=False)
        .rename("Frequency")
        .reset_index()
        .rename(columns={"partner_group": "Partner type"})
    )
    partner_group_table["Percentage of collaborating PBOs"] = partner_group_table["Frequency"].map(
        lambda value: pct(value / collaborating_pbos * 100 if collaborating_pbos else 0.0)
    )

    raw_partner_table = (
        collaborating.assign(partner_type_raw=collaborating["partner_type_raw"].replace("", "Missing partner type"))
        .groupby("partner_type_raw")["pbo_key"]
        .nunique()
        .sort_values(ascending=False)
        .rename("Frequency")
        .reset_index()
        .rename(columns={"partner_type_raw": "Raw partner type"})
    )
    raw_partner_table["Percentage of collaborating PBOs"] = raw_partner_table["Frequency"].map(
        lambda value: pct(value / collaborating_pbos * 100 if collaborating_pbos else 0.0)
    )

    collaboration_type_rows = []
    for column in COLLABORATION_COLUMNS:
        frequency = int(collaborating.loc[is_flagged(collaborating[column]), "pbo_key"].nunique())
        collaboration_type_rows.append(
            {
                "Collaboration type": COLLABORATION_LABELS[column],
                "Frequency": frequency,
                "Percentage of collaborating PBOs": pct(frequency / collaborating_pbos * 100 if collaborating_pbos else 0.0),
            }
        )
    collaboration_type_table = pd.DataFrame(collaboration_type_rows).sort_values(
        "Frequency", ascending=False
    ).reset_index(drop=True)

    matrix_rows = []
    for partner_group, group in collaborating.groupby("partner_group"):
        partner_pbos = int(group["pbo_key"].nunique())
        for column in COLLABORATION_COLUMNS:
            frequency = int(group.loc[is_flagged(group[column]), "pbo_key"].nunique())
            matrix_rows.append(
                {
                    "Partner type": partner_group,
                    "Partner-type PBOs": partner_pbos,
                    "Collaboration type": COLLABORATION_LABELS[column],
                    "Frequency": frequency,
                    "Pct within partner type value": frequency / partner_pbos * 100 if partner_pbos else 0.0,
                    "Pct within partner type": pct(frequency / partner_pbos * 100 if partner_pbos else 0.0),
                    "Pct of collaborating PBOs": pct(frequency / collaborating_pbos * 100 if collaborating_pbos else 0.0),
                }
            )
    partner_collaboration_long = pd.DataFrame(matrix_rows).sort_values(
        ["Partner-type PBOs", "Partner type", "Frequency"],
        ascending=[False, True, False],
    ).reset_index(drop=True)

    matrix_wide = partner_collaboration_long.pivot(
        index="Partner type",
        columns="Collaboration type",
        values="Frequency",
    ).fillna(0).astype(int)
    matrix_pct_wide = partner_collaboration_long.pivot(
        index="Partner type",
        columns="Collaboration type",
        values="Pct within partner type value",
    )

    top_partner_group = partner_group_table.iloc[0]
    top_collaboration_type = collaboration_type_table.iloc[0]
    donor_funding_from = partner_collaboration_long[
        (partner_collaboration_long["Partner type"] == "Donor Agencies")
        & (partner_collaboration_long["Collaboration type"] == "Funding from partner")
    ].iloc[0]

    return {
        "raw_df": df,
        "coverage_table": coverage_table,
        "partner_group_table": partner_group_table,
        "raw_partner_table": raw_partner_table,
        "collaboration_type_table": collaboration_type_table,
        "partner_collaboration_long": partner_collaboration_long,
        "matrix_wide": matrix_wide,
        "matrix_pct_wide": matrix_pct_wide,
        "total_pbos": total_pbos,
        "collaborating_pbos": collaborating_pbos,
        "no_recorded_collaboration": no_recorded_collaboration,
        "top_partner_group": top_partner_group,
        "top_collaboration_type": top_collaboration_type,
        "donor_funding_from": donor_funding_from,
    }


def create_figures(data: dict[str, pd.DataFrame | int | float]) -> dict[str, go.Figure]:
    coverage_table = data["coverage_table"]
    partner_group_table = data["partner_group_table"]
    collaboration_type_table = data["collaboration_type_table"]
    matrix_wide = data["matrix_wide"]
    matrix_pct_wide = data["matrix_pct_wide"]

    coverage_chart = px.pie(
        coverage_table.iloc[:2],
        names="Status",
        values="PBOs",
        hole=0.58,
        color="Status",
        color_discrete_sequence=[PLOTLY_COLORS[1], PLOTLY_COLORS[4]],
        title="Collaboration Coverage Across PBOs",
    )
    coverage_chart.update_traces(textposition="inside", textinfo="percent+label")
    coverage_chart.update_layout(
        template="plotly_white",
        font=dict(size=14),
        margin=dict(l=30, r=30, t=70, b=30),
    )

    partner_chart = px.bar(
        partner_group_table,
        x="Frequency",
        y="Partner type",
        orientation="h",
        text="Percentage of collaborating PBOs",
        title="Partner Types Mentioned by Collaborating PBOs",
        color="Frequency",
        color_continuous_scale=["#BFD7EA", "#087E8B", "#0B3954"],
    )
    partner_chart.update_layout(
        template="plotly_white",
        yaxis=dict(categoryorder="total ascending"),
        coloraxis_showscale=False,
        font=dict(size=13),
        margin=dict(l=30, r=30, t=70, b=30),
    )
    partner_chart.update_traces(textposition="outside", hovertemplate="%{y}<br>PBOs=%{x}<br>%{text}<extra></extra>")

    collaboration_type_chart = px.bar(
        collaboration_type_table,
        x="Frequency",
        y="Collaboration type",
        orientation="h",
        text="Percentage of collaborating PBOs",
        title="Kinds of Collaboration Reported",
        color="Frequency",
        color_continuous_scale=["#F4D35E", "#FF5A5F", "#C81D25"],
    )
    collaboration_type_chart.update_layout(
        template="plotly_white",
        yaxis=dict(categoryorder="total ascending"),
        coloraxis_showscale=False,
        font=dict(size=13),
        margin=dict(l=30, r=30, t=70, b=30),
    )
    collaboration_type_chart.update_traces(textposition="outside", hovertemplate="%{y}<br>PBOs=%{x}<br>%{text}<extra></extra>")

    heatmap_text = matrix_wide.astype(str).values
    heatmap = go.Figure(
        data=go.Heatmap(
            z=matrix_pct_wide.values,
            x=list(matrix_pct_wide.columns),
            y=list(matrix_pct_wide.index),
            colorscale=[
                [0.0, "#F4F1DE"],
                [0.35, "#9BC1BC"],
                [0.7, "#3D5A80"],
                [1.0, "#1D3557"],
            ],
            text=heatmap_text,
            texttemplate="%{text}",
            hovertemplate="%{y}<br>%{x}<br>%{z:.2f}% within partner type<extra></extra>",
            colorbar=dict(title="% within<br>partner type"),
        )
    )
    heatmap.update_layout(
        title="Partner Type by Collaboration Type Heatmap",
        template="plotly_white",
        font=dict(size=12),
        margin=dict(l=30, r=30, t=70, b=30),
        xaxis_title="Collaboration type",
        yaxis_title="Partner type",
    )

    return {
        "coverage_chart": coverage_chart,
        "partner_chart": partner_chart,
        "collaboration_type_chart": collaboration_type_chart,
        "heatmap": heatmap,
    }


def write_dashboard(figures: dict[str, go.Figure]):
    REPORTS_DIR.mkdir(exist_ok=True)
    html_parts = [
        "<html><head><meta charset='utf-8'><title>Collaboration Report</title></head>",
        "<body style='font-family:Arial, sans-serif; margin: 18px;'>",
        "<h1>Collaboration Analysis Dashboard</h1>",
        "<p>Interactive Plotly companion for the collaboration.docx report.</p>",
    ]
    first = True
    for name, figure in figures.items():
        html_parts.append(
            figure.to_html(
                full_html=False,
                include_plotlyjs="inline" if first else False,
                config={"displayModeBar": False, "responsive": True},
            )
        )
        first = False
    html_parts.append("</body></html>")
    OUTPUT_HTML.write_text("\n".join(html_parts), encoding="utf-8")


def render_figure_images(figures: dict[str, go.Figure]) -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    rendered: dict[str, Path] = {}

    for name, figure in figures.items():
        html_path = ASSETS_DIR / f"{name}.html"
        png_path = ASSETS_DIR / f"{name}.png"
        figure.write_html(
            html_path,
            include_plotlyjs=True,
            full_html=True,
            config={"displayModeBar": False, "responsive": False},
        )
        command = [
            "timeout",
            "25s",
            "google-chrome",
            "--headless",
            "--disable-gpu",
            "--hide-scrollbars",
            "--window-size=1600,960",
            "--virtual-time-budget=12000",
            f"--screenshot={png_path}",
            html_path.resolve().as_uri(),
        ]
        result = subprocess.run(command, capture_output=True, text=True)
        if not png_path.exists() or png_path.stat().st_size == 0:
            raise RuntimeError(
                f"Plotly chart screenshot failed for {name}: "
                f"returncode={result.returncode}, stderr={result.stderr.strip()}"
            )
        rendered[name] = png_path

    return rendered


def build_report(data: dict[str, pd.DataFrame | int | float], image_paths: dict[str, Path]):
    coverage_table = data["coverage_table"]
    partner_group_table = data["partner_group_table"]
    raw_partner_table = data["raw_partner_table"]
    collaboration_type_table = data["collaboration_type_table"]
    partner_collaboration_long = data["partner_collaboration_long"]
    total_pbos = data["total_pbos"]
    collaborating_pbos = data["collaborating_pbos"]
    no_recorded_collaboration = data["no_recorded_collaboration"]
    top_partner_group = data["top_partner_group"]
    top_collaboration_type = data["top_collaboration_type"]
    donor_funding_from = data["donor_funding_from"]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(document, "Collaboration Analysis Report")
    add_subtitle(document, f"Source workbook: {SOURCE_PATH.name}")
    add_subtitle(document, f"Unique PBOs analyzed: {total_pbos:,}")

    document.add_heading("Scope and Method", level=1)
    add_bullet(document, "The unit of analysis for the main questions is the unique PBO name across the collaboration workbook.")
    add_bullet(document, "A PBO is treated as collaborating when at least one of the collaboration-type fields is recorded: information exchange, technical support, funding, or equipment flows.")
    add_bullet(document, "Partner-type analysis is deduplicated at PBO level, so the same PBO is counted once per partner type even if it appears in multiple rows or reporting periods.")
    add_bullet(document, "Partner labels were standardized into major groups for visuals, while a full raw partner-type table is included later in the report.")

    document.add_heading("Question 1: How Many PBOs Collaborated?", level=1)
    document.add_paragraph(
        f"Out of {total_pbos:,} unique PBOs in the collaboration workbook, {collaborating_pbos:,} had at least one recorded collaboration type. "
        f"That is {pct(collaborating_pbos / total_pbos * 100 if total_pbos else 0.0)} of all PBOs in this file. "
        f"{no_recorded_collaboration:,} PBOs ({pct(no_recorded_collaboration / total_pbos * 100 if total_pbos else 0.0)}) had partner rows but no collaboration type marked."
    )
    add_dataframe_table(document, coverage_table)
    add_figure(document, image_paths.get("coverage_chart", Path()), "Figure 1. Collaboration coverage across all PBOs.")

    document.add_heading("Which Partners Did Collaborating PBOs Work With?", level=1)
    document.add_paragraph(
        f"The most common partner group was {top_partner_group['Partner type']}, mentioned by {int(top_partner_group['Frequency']):,} collaborating PBOs "
        f"({top_partner_group['Percentage of collaborating PBOs']}). The table and chart below show the frequency and percentage for each cleaned partner group."
    )
    add_dataframe_table(document, partner_group_table)
    add_figure(document, image_paths.get("partner_chart", Path()), "Figure 2. Partner groups mentioned by collaborating PBOs.")

    document.add_heading("Question 2: What Kind of Partnership Was It?", level=1)
    document.add_paragraph(
        f"The most common collaboration type was {top_collaboration_type['Collaboration type']}, reported by {int(top_collaboration_type['Frequency']):,} collaborating PBOs "
        f"({top_collaboration_type['Percentage of collaborating PBOs']})."
    )
    add_dataframe_table(document, collaboration_type_table)
    add_figure(document, image_paths.get("collaboration_type_chart", Path()), "Figure 3. Collaboration types reported by collaborating PBOs.")

    document.add_heading("Partner Type by Collaboration Type", level=1)
    document.add_paragraph(
        "The matrix below shows, for each cleaned partner group, how many unique PBOs reported each collaboration type. The heatmap visual summarizes the same pattern using percentages within each partner type."
    )
    add_dataframe_table(document, partner_collaboration_long)
    add_figure(document, image_paths.get("heatmap", Path()), "Figure 4. Heatmap of collaboration types within each partner group.")

    document.add_heading("Raw Partner Types Captured in the Workbook", level=1)
    document.add_paragraph(
        "This appendix-style table keeps the original partner-type labels as entered in the source workbook, so one-off labels and specific named partners remain visible for audit and cleaning purposes."
    )
    add_dataframe_table(document, raw_partner_table)

    document.add_page_break()
    document.add_heading("Headline Findings", level=1)
    document.add_paragraph(
        f"{collaborating_pbos:,} out of {total_pbos:,} PBOs in the collaboration workbook reported at least one collaboration type, indicating very broad collaboration coverage in this dataset."
    )
    document.add_paragraph(
        f"PBO-to-PBO collaboration was the most common partner grouping, followed by CBOs and Government Agencies."
    )
    document.add_paragraph(
        f"Information exchange dominated the collaboration modes overall, while technical support flows were also common in both directions."
    )
    document.add_paragraph(
        f"Among donor-agency collaborations, funding from partner was recorded by {int(donor_funding_from['Frequency']):,} unique PBOs, "
        f"which is {donor_funding_from['Pct within partner type']} within the Donor Agencies partner group."
    )

    document.save(OUTPUT_DOCX)


def main():
    data = prepare_data()
    figures = create_figures(data)
    write_dashboard(figures)
    image_paths = render_figure_images(figures)
    build_report(data, image_paths)
    print(OUTPUT_DOCX)
    print(OUTPUT_HTML)


if __name__ == "__main__":
    main()
