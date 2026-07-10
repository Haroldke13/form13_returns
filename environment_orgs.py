#!/usr/bin/env python3
"""Build a pbo_name/sector list for ENVIRONMENT organizations."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


DEFAULT_INPUT = Path(__file__).resolve().parent / "projectIMPLEMENTATIONS.xlsx"
DEFAULT_DOCX_OUTPUT = Path(__file__).resolve().parent / "environment_orgs.docx"
SHEET_NAME = "Matched Projects"
PBO_COLUMN = "pbo_name"
SECTOR_COLUMN = "project_sector"
TARGET_SECTOR = "ENVIRONMENT"


def build_environment_orgs(input_path: Path) -> list[dict[str, str]]:
    """Return unique pbo_name/sector pairs for ENVIRONMENT records."""
    data = pd.read_excel(
        input_path,
        sheet_name=SHEET_NAME,
        usecols=[PBO_COLUMN, SECTOR_COLUMN],
    ).copy()

    data[PBO_COLUMN] = data[PBO_COLUMN].fillna("").astype(str).str.strip()
    data[SECTOR_COLUMN] = (
        data[SECTOR_COLUMN].fillna("").astype(str).str.strip().str.upper()
    )

    filtered = (
        data.loc[
            data[PBO_COLUMN].ne("") & data[SECTOR_COLUMN].eq(TARGET_SECTOR),
            [PBO_COLUMN, SECTOR_COLUMN],
        ]
        .drop_duplicates()
        .sort_values([PBO_COLUMN, SECTOR_COLUMN], kind="stable")
        .rename(columns={SECTOR_COLUMN: "sector"})
    )

    return filtered.to_dict(orient="records")


environment_orgs: list[dict[str, str]] = []
if DEFAULT_INPUT.exists():
    environment_orgs = build_environment_orgs(DEFAULT_INPUT)


def write_docx(
    records: list[dict[str, str]], output_path: Path, source_path: Path
) -> Path:
    """Write the environment organization list to a DOCX table."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    document.add_heading("ENVIRONMENT Organizations", level=0)
    document.add_paragraph(f"Source workbook: {source_path.name}")
    document.add_paragraph(f"Unique organizations listed: {len(records)}")

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "PBO Name"
    header_cells[1].text = "Sector"

    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = record["pbo_name"]
        row_cells[1].text = record["sector"]

    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Create a pbo_name/sector list for organizations dealing with "
            "ENVIRONMENT from projectIMPLEMENTATIONS.xlsx."
        )
    )
    parser.add_argument(
        "--input",
        type=Path,
        default=DEFAULT_INPUT,
        help=f"Workbook to read (default: {DEFAULT_INPUT})",
    )
    parser.add_argument(
        "--docx-output",
        nargs="?",
        const=DEFAULT_DOCX_OUTPUT,
        type=Path,
        help=(
            "Write the list to a DOCX file. If no path is supplied, "
            f"{DEFAULT_DOCX_OUTPUT} is used."
        ),
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    selected_orgs = (
        environment_orgs
        if args.input.resolve() == DEFAULT_INPUT.resolve() and environment_orgs
        else build_environment_orgs(args.input)
    )
    if args.docx_output:
        output_path = args.docx_output
        write_docx(selected_orgs, output_path, args.input)
        print(output_path)
        return

    print(json.dumps(selected_orgs, indent=2))


if __name__ == "__main__":
    main()
